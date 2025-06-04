import os
import zipfile
from docx2pdf import convert
from lxml import etree
from PIL import Image
from io import BytesIO
from docx import Document
from OpinionAnalysis.settings import BASE_DIR
from doc_project.common_utils import save_html_file, save_json_file
from doc_project.generater import structured_data_to_html


def extract_images_from_drawing(drawing_element, doc, image_output_dir):
    """
    从 <w:drawing> 元素提取图片，保存并返回图片路径
    """
    # 命名空间映射
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    blips = drawing_element.findall('.//a:blip', namespaces=nsmap)
    for blip in blips:
        embed_rid = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not embed_rid:
            continue
        image_part = doc.part.related_parts.get(embed_rid)
        if image_part:
            image_bytes = image_part.blob
            image = Image.open(BytesIO(image_bytes))
            # 图片保存路径
            os.makedirs(image_output_dir, exist_ok=True)
            image_filename = f'image_{embed_rid}.png'
            image_path = os.path.join(image_output_dir, image_filename)
            image.save(image_path)
            return 'docx_images/'+image_filename
    return None


"""
返回的结构化数据示例：
{'type': 'paragraph', 'style': 'Normal', 'text': '正文内容示例', 'font_size': 12}
{'type': 'table', 'data': [['单元格11', '单元格12'], ['单元格21', '单元格22']]}
{'type': 'image', 'path': 'extracted_images/uuid.png'}
"""
def analyze_docx_structured(docx_path, image_output_dir):
    doc = Document(docx_path)
    body = doc.element.body
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    structured_data = []

    para_idx = 0
    table_idx = 0

    for child in body.iterchildren():
        tag = etree.QName(child).localname
        if tag == 'p':
            para = doc.paragraphs[para_idx]
            para_idx += 1

            style = para.style.name if para.style else 'Normal'
            text = para.text.strip()

            # 提取段落中的图片
            drawing_elements = child.findall('.//w:drawing', namespaces=ns)
            for drawing in drawing_elements:
                img_path = extract_images_from_drawing(drawing, doc, image_output_dir)
                if img_path:
                    structured_data.append({'type': 'image', 'path': img_path})

            if text:
                # 尝试获取第一个 run 的字号，单位是 Pt
                font_size = None
                if para.runs:
                    first_run_size = para.runs[0].font.size
                    if first_run_size is not None:
                        font_size = first_run_size.pt  # 转成浮点磅数
                structured_data.append({'type': 'paragraph', 'style': style, 'text': text, 'font_size': font_size})

        elif tag == 'tbl':
            table = doc.tables[table_idx]
            table_idx += 1

            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)

            structured_data.append({'type': 'table', 'data': table_data})

    return structured_data

uploads_dir = os.path.join(BASE_DIR,'static', 'uploads')
downloads_dir = os.path.join(BASE_DIR,'static', 'downloads')
html_path = os.path.join(downloads_dir, 'docx_html.html')
image_output_dir = os.path.join(downloads_dir, 'docx_images')
json_path = os.path.join(downloads_dir, 'docx_json.json')
pdf_path = os.path.join(downloads_dir, 'docx2pdf.pdf')

transformed_path= 'transformed_docx.zip'
def transform_docx(file_name):
    docx_path = os.path.join(uploads_dir, file_name)

    # 结构化数据
    structured_data = analyze_docx_structured(docx_path, image_output_dir)
    # 保存 json 文件
    save_json_file("docx_json.json", structured_data)
    # 生成html
    html_str = structured_data_to_html(structured_data)
    # 保存 html 文件
    save_html_file("docx_html.html", html_str)
    # 转换为 PDF
    docx_to_pdf(docx_path)

    # 创建 zip 包
    zip_path = os.path.join(downloads_dir, transformed_path)
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        # 添加 HTML 文件
        zipf.write(html_path, arcname='docx_html.html')
        zipf.write(json_path, arcname='docx_json.json')
        zipf.write(pdf_path, arcname='docx_pdf.pdf')
        # 添加图片资源
        for root, _, files in os.walk(image_output_dir):
            for f in files:
                img_full_path = os.path.join(root, f)
                rel_path = os.path.relpath(img_full_path, downloads_dir)
                zipf.write(img_full_path, arcname=rel_path)

    max_elements = 10
    preview_html_str =  structured_data_to_html(structured_data[:max_elements],is_preview=True)
    return {
        'output_file': transformed_path,
        'preview': preview_html_str,
    }


def docx_to_pdf(docx_path):
    convert(docx_path,pdf_path)

def docx_to_xml(docx_path):
    doc = Document(docx_path)
    # 访问底层 XML Element 对象
    xml_element = doc.element
    # 将 XML 对象转成字符串
    xml_str = xml_element.xml
    return xml_str





if __name__ == '__main__':
    pass
    # docx_path = os.path.join(uploads_dir, 'example.docx')
    # image_output_dir = os.path.join(downloads_dir, 'docx_images')
    # # print(docx_to_xml(docx_path))
    # structured_data = analyze_docx_structured(docx_path, image_output_dir)
    # print(structured_data)